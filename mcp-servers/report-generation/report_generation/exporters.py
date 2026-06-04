"""Export helpers for rendering reports to PDF/Word and uploading to S3.

Heavy / native dependencies (``weasyprint``, ``docx``, ``boto3``) are imported
lazily inside each function so that this module — and the rest of the package —
can be imported and unit-tested without those libraries installed.

Requirements: 5.4, 5.5
"""

from __future__ import annotations

from .models import ReportContent


def export_pdf(report: ReportContent, include_charts: bool = True) -> bytes:
    """Render the report to a PDF document using WeasyPrint.

    Args:
        report: The structured report content.
        include_charts: Whether to embed inline charts (Requirement 5.4).

    Returns:
        The PDF file contents as bytes.

    Raises:
        RuntimeError: If WeasyPrint (or its native dependencies) is unavailable.
    """
    try:
        from weasyprint import HTML  # lazy import: native libs required
    except Exception as exc:  # pragma: no cover - depends on native libs
        raise RuntimeError(
            "WeasyPrint is not available. Install 'weasyprint' and its native "
            "dependencies to enable PDF export."
        ) from exc

    from .templates import render_html  # lazy import: keeps jinja2 optional at import time

    html = render_html(report, include_charts=include_charts)
    return HTML(string=html).write_pdf()


def export_docx(report: ReportContent, include_charts: bool = True) -> bytes:
    """Render the report to a Word (.docx) document using python-docx.

    Args:
        report: The structured report content.
        include_charts: Whether to embed inline charts (Requirement 5.4).

    Returns:
        The .docx file contents as bytes.

    Raises:
        RuntimeError: If python-docx is unavailable.
    """
    try:
        import base64
        import io

        from docx import Document  # lazy import
        from docx.shared import Pt
    except Exception as exc:  # pragma: no cover - depends on optional lib
        raise RuntimeError(
            "python-docx is not available. Install 'python-docx' to enable Word export."
        ) from exc

    document = Document()

    # Title.
    document.add_heading(report.title, level=0)

    # Data-source metadata header (Requirement 5.7).
    meta = report.metadata
    document.add_heading("数据来源信息 (Data Source Metadata)", level=2)
    meta_table = document.add_table(rows=4, cols=2)
    meta_rows = [
        ("数据文件 (File Name)", meta.file_name or "未提供"),
        ("上传时间 (Upload Time)", meta.upload_time or "未提供"),
        ("数据行数 (Row Count)", str(meta.row_count)),
        ("数据列数 (Column Count)", str(meta.column_count)),
    ]
    for row, (label, value) in zip(meta_table.rows, meta_rows):
        row.cells[0].text = label
        row.cells[1].text = value

    # Sections.
    for section in report.sections:
        document.add_heading(section.title, level=1)
        _add_body_to_docx(document, section.body, Pt)

        if include_charts:
            for chart in section.charts:
                if chart.image_data_uri and _is_embeddable_data_uri(chart.image_data_uri):
                    try:
                        raw = base64.b64decode(chart.image_data_uri.split(",", 1)[1])
                        document.add_picture(io.BytesIO(raw))
                    except Exception:
                        document.add_paragraph(f"[图表: {chart.title}]")
                else:
                    document.add_paragraph(f"[图表: {chart.title}]")
                caption = chart.title + (f" — {chart.caption}" if chart.caption else "")
                document.add_paragraph(caption).italic = True

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _add_body_to_docx(document, body: str, pt) -> None:
    """Append a section body string to the document as paragraphs/bullets."""
    if not body:
        return
    for raw_line in body.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith(("- ", "* ")):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            document.add_paragraph(line)


def _is_embeddable_data_uri(value: str) -> bool:
    """Return True for base64 image data URIs that python-docx can embed."""
    return value.startswith("data:image/") and ";base64," in value


def upload_to_s3(
    data: bytes,
    bucket: str,
    key: str,
    content_type: str | None = None,
    expires_in: int = 3600,
) -> str:
    """Upload bytes to S3 and return a download URL.

    Returns a presigned GET URL valid for ``expires_in`` seconds when possible,
    falling back to the canonical ``s3://`` URI if presigning fails.

    Args:
        data: File contents to upload.
        bucket: Target S3 bucket name.
        key: Object key within the bucket.
        content_type: Optional MIME type for the object.
        expires_in: Presigned URL lifetime in seconds.

    Returns:
        A download URL (presigned HTTPS URL or ``s3://`` URI).

    Raises:
        RuntimeError: If boto3 is unavailable.
    """
    try:
        import boto3  # lazy import
    except Exception as exc:  # pragma: no cover - optional dep
        raise RuntimeError("boto3 is not available. Install 'boto3' to enable S3 upload.") from exc

    client = boto3.client("s3")
    put_kwargs: dict = {"Bucket": bucket, "Key": key, "Body": data}
    if content_type:
        put_kwargs["ContentType"] = content_type
    client.put_object(**put_kwargs)

    try:
        return client.generate_presigned_url(
            "get_object",
            Params={"Bucket": bucket, "Key": key},
            ExpiresIn=expires_in,
        )
    except Exception:  # pragma: no cover - depends on credentials/config
        return f"s3://{bucket}/{key}"


# Format -> (exporter, content_type, extension) registry.
EXPORT_FORMATS: dict[str, tuple] = {
    "pdf": (export_pdf, "application/pdf", "pdf"),
    "docx": (
        export_docx,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "docx",
    ),
}


def export_report_bytes(report: ReportContent, fmt: str = "pdf", include_charts: bool = True) -> tuple[bytes, str, str]:
    """Export a report to the requested format.

    Args:
        report: The structured report content.
        fmt: Output format, one of ``"pdf"`` or ``"docx"``.
        include_charts: Whether to embed inline charts.

    Returns:
        Tuple of ``(data_bytes, content_type, file_extension)``.

    Raises:
        ValueError: If ``fmt`` is not a supported format.
    """
    normalized = fmt.lower().lstrip(".")
    if normalized not in EXPORT_FORMATS:
        raise ValueError(f"Unsupported export format: {fmt!r}. Supported: {sorted(EXPORT_FORMATS)}")
    exporter, content_type, extension = EXPORT_FORMATS[normalized]
    data = exporter(report, include_charts=include_charts)
    return data, content_type, extension
